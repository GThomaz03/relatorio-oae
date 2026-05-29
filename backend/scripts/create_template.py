"""Gera report_template.docx a partir do modelo RSP de referência."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "report_template.docx"
DEFAULT_REFERENCE = Path(r"d:\RSP-116RJ-233-243-ACA-FUN-RT-L1-001-R00 (1).docx")


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _find_paragraph_index(doc: Document, predicate) -> int | None:
    for i, para in enumerate(doc.paragraphs):
        if predicate(para):
            return i
    return None


def _add_paragraph(doc: Document, text: str, style: str) -> None:
    doc.add_paragraph(text, style=style)


def create_template(
    path: Path | None = None,
    reference_path: Path | None = None,
) -> Path:
    """
    Copia o DOCX de referência e injeta placeholders docxtpl nas secções dinâmicas.

    Ordem final:
    1. ANOMALIAS CONSTATADAS (Heading 2, estático)
    2. Loops de macro/elemento/linha
    3. ANEXO VI — RELATÓRIO FOTOGRÁFICO (Title, estático)
    4. Loop de fotografias
    """
    target = path or TEMPLATE_PATH
    reference = reference_path or DEFAULT_REFERENCE
    if not reference.is_file():
        raise FileNotFoundError(f"Documento de referência não encontrado: {reference}")

    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, target)

    doc = Document(str(target))

    annex_idx = _find_paragraph_index(
        doc,
        lambda p: "ANEXO" in p.text.upper() and "FOTOGR" in p.text.upper(),
    )
    anomaly_idx = _find_paragraph_index(
        doc,
        lambda p: "ANOMALIAS CONSTATADAS" in p.text.upper(),
    )

    if annex_idx is None:
        raise ValueError("Secção 'ANEXO ... RELATÓRIO FOTOGRÁFICO' não encontrada no modelo.")
    if anomaly_idx is None:
        raise ValueError("Secção 'ANOMALIAS CONSTATADAS' não encontrada no modelo.")

    annex_title = doc.paragraphs[annex_idx].text
    annex_style = doc.paragraphs[annex_idx].style.name

    # Remove todo o conteúdo após o título ANOMALIAS CONSTATADAS
    while len(doc.paragraphs) > anomaly_idx + 1:
        _delete_paragraph(doc.paragraphs[-1])

    # --- Secção de anomalias (placeholders) ---
    _add_paragraph(doc, "{% for macro in anomaly_macros %}", "Normal")
    _add_paragraph(doc, "{{ macro.title }}", "Heading 4")
    _add_paragraph(doc, "{% for element in macro.elements %}", "Normal")
    _add_paragraph(doc, "{% if element.title %}{{ element.title }}{% endif %}", "Heading 4")
    _add_paragraph(doc, "{% for line in element.lines %}", "Normal")
    _add_paragraph(doc, "{{ line.line }}", "Normal")
    _add_paragraph(doc, "{% endfor %}", "Normal")
    _add_paragraph(doc, "{% endfor %}", "Normal")
    _add_paragraph(doc, "{% endfor %}", "Normal")

    # Quebra de página antes do anexo fotográfico
    page_break_para = doc.add_paragraph()
    page_break_para.add_run().add_break(WD_BREAK.PAGE)

    # --- Anexo fotográfico ---
    _add_paragraph(doc, annex_title, annex_style)
    _add_paragraph(doc, "{% for photo in photos %}", "Normal")
    _add_paragraph(doc, "{{ photo.image }}", "Normal")
    _add_paragraph(doc, "{{ photo.code }}", "Caption")
    _add_paragraph(doc, "{{ photo.location_line }}", "Normal")
    _add_paragraph(doc, "{{ photo.description_line }}", "Normal")
    _add_paragraph(doc, "{% endfor %}", "Normal")

    doc.save(str(target))
    return target


if __name__ == "__main__":
    out = create_template()
    print(f"Template criado: {out}")
