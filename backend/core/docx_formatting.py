"""Pós-processamento de formatação do documento Word gerado."""

from __future__ import annotations

import logging
import re

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from backend.config import (
    PHOTO_BLOCK_PARAGRAPH_COUNT,
    PHOTO_PAGE_LEADING_BLANK_LINES,
    PHOTO_PAGE_MIDDLE_BLANK_LINES,
    PHOTOS_PER_PAGE,
)

logger = logging.getLogger(__name__)

# Estilos de título — preservam formatação original do template RSP
HEADING_STYLE_PREFIXES: tuple[str, ...] = (
    "Title",
    "Heading",
    "TOC",
)

BODY_FONT_NAME = "Arial"
BODY_FONT_SIZE_PT = 10

_ANNEX_PHOTO_PATTERN = re.compile(r"ANEXO.*FOTOGR", re.IGNORECASE)
_DRAWING_TAG = qn("w:drawing")
_PICT_TAG = qn("w:pict")
_OBJECT_TAG = qn("w:object")


def _is_heading_style(style_name: str) -> bool:
    return any(style_name.startswith(prefix) for prefix in HEADING_STYLE_PREFIXES)


def _is_annex_photo_title(text: str) -> bool:
    return bool(_ANNEX_PHOTO_PATTERN.search(text))


def _paragraph_has_inline_image(paragraph: Paragraph) -> bool:
    element = paragraph._element
    return bool(
        element.findall(f".//{_DRAWING_TAG}")
        or element.findall(f".//{_PICT_TAG}")
        or element.findall(f".//{_OBJECT_TAG}")
    )


def _apply_arial_10_to_paragraph(paragraph: Paragraph) -> None:
    """Aplica Arial 10 a todos os runs do parágrafo (corpo/legendas)."""
    for run in paragraph.runs:
        run.font.name = BODY_FONT_NAME
        run.font.size = Pt(BODY_FONT_SIZE_PT)
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = run._element.makeelement(qn("w:rFonts"), {})
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), BODY_FONT_NAME)
        r_fonts.set(qn("w:hAnsi"), BODY_FONT_NAME)
        r_fonts.set(qn("w:cs"), BODY_FONT_NAME)


def _insert_page_break_before(paragraph: Paragraph) -> None:
    """Insere quebra de página imediatamente antes do parágrafo indicado."""
    new_p = paragraph.insert_paragraph_before("")
    new_p.add_run().add_break(WD_BREAK.PAGE)


def _is_removable_empty_paragraph(paragraph: Paragraph) -> bool:
    if paragraph.text.strip():
        return False
    if _paragraph_has_inline_image(paragraph):
        return False
    if 'w:type="page"' in paragraph._element.xml:
        return False
    return True


def _insert_blank_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Insere parágrafo vazio centralizado imediatamente após o parágrafo indicado."""
    new_p_el = OxmlElement("w:p")
    paragraph._element.addnext(new_p_el)
    new_p = Paragraph(new_p_el, paragraph._parent)
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_arial_10_to_paragraph(new_p)
    return new_p


def _ensure_blank_lines_before_index(document: Document, target_idx: int, required: int) -> None:
    """Garante exatamente `required` parágrafos vazios consecutivos antes do alvo (por índice)."""
    for _ in range(20):
        paragraphs = list(document.paragraphs)
        if target_idx >= len(paragraphs):
            return

        blank_count = 0
        scan_idx = target_idx - 1
        while scan_idx >= 0 and _is_removable_empty_paragraph(paragraphs[scan_idx]):
            blank_count += 1
            scan_idx -= 1

        if blank_count > required:
            paragraphs[target_idx - 1]._element.getparent().remove(paragraphs[target_idx - 1]._element)
            target_idx -= 1
            continue
        if blank_count < required:
            if target_idx <= 0:
                return
            _insert_blank_paragraph_after(paragraphs[target_idx - 1])
            target_idx += 1
            continue
        return


def _ensure_blank_lines_before(paragraph: Paragraph, required: int) -> None:
    """Garante exatamente `required` parágrafos vazios consecutivos antes do alvo."""
    document = paragraph.part.document
    paragraphs = list(document.paragraphs)
    try:
        target_idx = paragraphs.index(paragraph)
    except ValueError:
        return
    _ensure_blank_lines_before_index(document, target_idx, required)


def _add_page_break_after(paragraph: Paragraph) -> None:
    """Insere quebra de página ao final do parágrafo indicado."""
    if paragraph.runs:
        paragraph.runs[-1].add_break(WD_BREAK.PAGE)
    else:
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def _find_annex_title_index(paragraphs: list[Paragraph]) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if _is_annex_photo_title(paragraph.text):
            return idx
    return None


def _annex_image_indices(document: Document) -> tuple[int | None, list[int]]:
    paragraphs = list(document.paragraphs)
    annex_idx = _find_annex_title_index(paragraphs)
    if annex_idx is None:
        return None, []
    image_indices = [
        idx
        for idx, paragraph in enumerate(paragraphs)
        if idx > annex_idx and _paragraph_has_inline_image(paragraph)
    ]
    return annex_idx, image_indices


def _block_end_index(image_start: int, next_image_start: int | None, paragraph_count: int) -> int:
    default_end = min(image_start + PHOTO_BLOCK_PARAGRAPH_COUNT, paragraph_count)
    if next_image_start is None:
        return default_end - 1
    return min(default_end, next_image_start) - 1


def apply_photo_annex_page_layout(document: Document) -> None:
    """
    Aplica layout fixo do anexo fotográfico:

    - 2 fotos por página (última página pode ter 1)
    - 2 linhas em branco antes da 1.ª foto de cada página
    - 4 linhas em branco entre a legenda da 1.ª e a 2.ª foto
    - quebra de página após a 2.ª foto da página
    """
    annex_idx, image_indices = _annex_image_indices(document)
    if annex_idx is None or not image_indices:
        return

    pages = [
        image_indices[i : i + PHOTOS_PER_PAGE]
        for i in range(0, len(image_indices), PHOTOS_PER_PAGE)
    ]

    for page_num in reversed(range(len(pages))):
        _, current_images = _annex_image_indices(document)
        page_images = current_images[page_num * PHOTOS_PER_PAGE : page_num * PHOTOS_PER_PAGE + len(pages[page_num])]
        paragraphs = list(document.paragraphs)

        if len(page_images) == PHOTOS_PER_PAGE:
            _ensure_blank_lines_before_index(document, page_images[1], PHOTO_PAGE_MIDDLE_BLANK_LINES)

            _, current_images = _annex_image_indices(document)
            page_images = current_images[page_num * PHOTOS_PER_PAGE : page_num * PHOTOS_PER_PAGE + PHOTOS_PER_PAGE]
            paragraphs = list(document.paragraphs)
            next_image = (
                current_images[page_num * PHOTOS_PER_PAGE + PHOTOS_PER_PAGE]
                if page_num * PHOTOS_PER_PAGE + PHOTOS_PER_PAGE < len(current_images)
                else None
            )
            block_end = _block_end_index(page_images[1], next_image, len(paragraphs))
            _add_page_break_after(paragraphs[block_end])

        _, current_images = _annex_image_indices(document)
        page_images = current_images[page_num * PHOTOS_PER_PAGE : page_num * PHOTOS_PER_PAGE + len(pages[page_num])]
        _ensure_blank_lines_before_index(document, page_images[0], PHOTO_PAGE_LEADING_BLANK_LINES)

    logger.info(
        "Layout fotográfico: %d foto(s) em %d página(s)",
        len(image_indices),
        len(pages),
    )


def apply_report_formatting(document: Document) -> None:
    """
    Aplica formatação final ao relatório renderizado:

    - Quebra de página antes do ANEXO fotográfico
    - Corpo e legendas em Arial 10 (títulos preservados)
    - Secção fotográfica centralizada
    - Layout 2 fotos por página com espaçamento padronizado
    """
    paragraphs = list(document.paragraphs)
    annex_start_index: int | None = None

    for idx, paragraph in enumerate(paragraphs):
        if _is_annex_photo_title(paragraph.text):
            annex_start_index = idx
            break

    if annex_start_index is not None and annex_start_index > 0:
        prev_para = paragraphs[annex_start_index - 1]
        prev_xml = prev_para._element.xml
        if 'w:type="page"' not in prev_xml:
            _insert_page_break_before(paragraphs[annex_start_index])
            paragraphs = list(document.paragraphs)
            for idx, paragraph in enumerate(paragraphs):
                if _is_annex_photo_title(paragraph.text):
                    annex_start_index = idx
                    break

    in_annex = False
    for paragraph in document.paragraphs:
        if _is_annex_photo_title(paragraph.text):
            in_annex = True

        style_name = paragraph.style.name if paragraph.style else "Normal"

        if not _is_heading_style(style_name):
            _apply_arial_10_to_paragraph(paragraph)

        if in_annex:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    apply_photo_annex_page_layout(document)

    logger.info(
        "Formatação aplicada: Arial 10 no corpo, anexo centralizado%s",
        ", quebra de página inserida" if annex_start_index is not None else "",
    )
