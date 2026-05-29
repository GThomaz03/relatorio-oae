"""Testes de formatação pós-render do DOCX."""

from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from backend.config import (
    PHOTO_HEIGHT_MM,
    PHOTO_PAGE_LEADING_BLANK_LINES,
    PHOTO_PAGE_MIDDLE_BLANK_LINES,
    PHOTO_WIDTH_MM,
)
from docx.oxml.ns import qn
from backend.core.docx_formatting import apply_report_formatting
from backend.core.parser_excel import parse_excel
from backend.core.parser_images import attach_images_to_anomalies
from backend.core.photo_generator import build_photo_report
from backend.core.text_generator import build_sections
from backend.core.word_generator import render_report
from backend.models.inspection import InspectionReport, ReportMetadata

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}


def _read_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def test_apply_report_formatting_arial_and_center(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("ANOMALIAS CONSTATADAS", style="Heading 2")
    doc.add_paragraph("Texto de anomalia", style="Normal")
    doc.add_paragraph("ANEXO VI — RELATÓRIO FOTOGRÁFICO", style="Title")
    doc.add_paragraph("IMG_0001", style="Caption")
    doc.add_paragraph("Legenda foto", style="Normal")

    apply_report_formatting(doc)

    body_para = doc.paragraphs[1]
    assert body_para.runs[0].font.name == "Arial"
    assert body_para.runs[0].font.size == Pt(10)

    annex_para = doc.paragraphs[3]
    assert annex_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def _w_tag(local: str) -> str:
    return f"{{{NS['w']}}}{local}"


def test_render_report_formatting(sample_excel, sample_images, sample_template, output_dir) -> None:
    anomalies, _ = parse_excel(sample_excel)
    anomalies, _ = attach_images_to_anomalies(anomalies, sample_images)
    groups = build_sections(anomalies)
    photo_report = build_photo_report(anomalies, groups=groups)

    report = InspectionReport(
        metadata=ReportMetadata(
            excel_path=sample_excel,
            images_dir=sample_images,
            template_path=sample_template,
            bridge_id="E116",
        ),
        anomalies=anomalies,
        groups=groups,
        photo_report=photo_report,
    )

    out = output_dir / "formatting_test.docx"
    render_report(report, out, template_path=sample_template)

    xml = _read_document_xml(out)
    root = ET.fromstring(xml)

    # Quebra de página presente
    page_breaks = [
        el
        for el in root.iter(_w_tag("br"))
        if el.get(_w_tag("type")) == "page"
    ]
    assert len(page_breaks) >= 1

    # Arial 10 em parágrafo Normal (w:sz val=20 => 10pt)
    sizes = [el.get(_w_tag("val")) for el in root.iter(_w_tag("sz"))]
    assert "20" in sizes

    # Alinhamento centro no anexo
    jc_values = [el.get(_w_tag("val")) for el in root.iter(_w_tag("jc"))]
    assert "center" in jc_values

    # Dimensões da imagem inline (EMU: 1 mm = 36000 EMU)
    expected_cx = str(int(PHOTO_WIDTH_MM * 36000))
    expected_cy = str(int(PHOTO_HEIGHT_MM * 36000))
    extents = root.findall(".//wp:extent", NS)
    assert extents, "Nenhuma imagem inline encontrada no DOCX"
    assert any(ext.get("cx") == expected_cx and ext.get("cy") == expected_cy for ext in extents)


def test_render_report_photo_page_layout(sample_excel, sample_images, sample_template, output_dir) -> None:
    """Anexo com 3 fotos: página 1 (2 fotos + quebra), página 2 (1 foto, sem quebra final)."""
    anomalies, _ = parse_excel(sample_excel)
    anomalies, _ = attach_images_to_anomalies(anomalies, sample_images)
    groups = build_sections(anomalies)
    photo_report = build_photo_report(anomalies, groups=groups)

    report = InspectionReport(
        metadata=ReportMetadata(
            excel_path=sample_excel,
            images_dir=sample_images,
            template_path=sample_template,
            bridge_id="E116",
        ),
        anomalies=anomalies,
        groups=groups,
        photo_report=photo_report,
    )

    out = output_dir / "photo_layout_test.docx"
    render_report(report, out, template_path=sample_template)

    doc = Document(str(out))
    annex_idx = next(
        i for i, p in enumerate(doc.paragraphs) if "ANEXO" in p.text.upper() and "FOTOGR" in p.text.upper()
    )

    image_indices = [
        i
        for i, p in enumerate(doc.paragraphs)
        if i > annex_idx and p._element.findall(f".//{qn('w:drawing')}")
    ]
    assert len(image_indices) == 3

    first_image_idx = image_indices[0]
    blanks_before_first = 0
    for i in range(first_image_idx - 1, annex_idx, -1):
        if not doc.paragraphs[i].text.strip():
            blanks_before_first += 1
        else:
            break
    assert blanks_before_first == PHOTO_PAGE_LEADING_BLANK_LINES

    second_image_idx = image_indices[1]
    blanks_between = 0
    for i in range(second_image_idx - 1, first_image_idx, -1):
        if not doc.paragraphs[i].text.strip():
            blanks_between += 1
        else:
            break
    assert blanks_between == PHOTO_PAGE_MIDDLE_BLANK_LINES

    xml = _read_document_xml(out)
    page_breaks = [
        el
        for el in ET.fromstring(xml).iter(_w_tag("br"))
        if el.get(_w_tag("type")) == "page"
    ]
    assert len(page_breaks) >= 2


def test_default_output_filename_includes_time(sample_excel, sample_images, sample_template) -> None:
    from backend.core.word_generator import default_output_filename

    report = InspectionReport(
        metadata=ReportMetadata(
            excel_path=sample_excel,
            images_dir=sample_images,
            template_path=sample_template,
            bridge_id="E116",
        ),
    )
    name = default_output_filename(report)
    assert name.endswith("_relatorio.docx")
    parts = name.replace("_relatorio.docx", "").split("_")
    assert len(parts) >= 3  # E116, YYYYMMDD, HHMMSS
